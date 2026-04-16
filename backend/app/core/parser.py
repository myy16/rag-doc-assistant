import subprocess
from pathlib import Path


def parse_document(file_path: str, file_type: str) -> str:
    """
    Parse a document file and return its plain text content.

    Args:
        file_path: Absolute or relative path to the saved file.
        file_type: Extension without dot — 'pdf', 'docx', 'doc', 'txt'.

    Returns:
        Extracted plain text as a string.

    Raises:
        ValueError: If file_type is not supported.
        RuntimeError: If parsing fails.
    """
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "doc": _parse_doc,
        "txt": _parse_txt,
    }

    if file_type not in parsers:
        raise ValueError(f"Unsupported file type: '{file_type}'")

    return parsers[file_type](file_path)


def _parse_pdf(file_path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber is not installed. Run: pip install pdfplumber")

    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF '{file_path}': {e}")


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    try:
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX '{file_path}': {e}")


def _parse_doc(file_path: str) -> str:
    # First try antiword (real binary .doc format)
    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except FileNotFoundError:
        pass

    # Fallback: file may actually be DOCX saved with .doc extension
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text:
            return text
    except Exception:
        pass

    raise RuntimeError(
        f"Failed to parse DOC '{file_path}': file is not a valid .doc or .docx format."
    )


def _parse_txt(file_path: str) -> str:
    path = Path(file_path)
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1254"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    raise RuntimeError(
        f"Failed to decode TXT file '{file_path}' with supported encodings."
    )
